"""
Export anti-strat report to a Word (.docx) document.

Usage:
  python export_docx.py --team "G2 Esports" --series 103891 --map Pearl
"""

import argparse
import io
import sys
import subprocess
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: Run: pip install python-docx")
    sys.exit(1)


# ── colour palette ────────────────────────────────────────────────────────────
RED    = RGBColor(180,  20,  20)
BLUE   = RGBColor( 31,  73, 125)
GOLD   = RGBColor(160, 110,   0)
GREEN  = RGBColor(  0, 120,  50)
GRAY   = RGBColor(110, 110, 110)
BLACK  = RGBColor( 20,  20,  20)
WHITE  = RGBColor(255, 255, 255)


def _shd(cell, fill_hex: str):
    """Set table-cell background."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)


def _p(doc, text='', bold=False, italic=False,
        color=BLACK, size=10, indent=0.0, align=None, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
    return p


def _h(doc, text, level=2, color=BLUE, size=None):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(3)
    return p


def _rule(doc):
    p = doc.add_paragraph('─' * 74)
    p.runs[0].font.size = Pt(6)
    p.runs[0].font.color.rgb = GRAY
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)


# ── section parsers ───────────────────────────────────────────────────────────

def _parse_sections(text: str) -> dict:
    """Split the captured CLI output into named sections."""
    sections = {
        'header':   [],
        'pool':     [],
        'roster':   [],
        'weapons':  [],
        'timing':   [],
        'sites':    [],
        'routes':   [],
        'rounds':   [],
        'brief':    [],
    }

    lines = text.splitlines()
    cur = 'header'
    i = 0
    while i < len(lines):
        L = lines[i]
        s = L.strip()

        if 'MAP POOL' in s and s.startswith('MAP POOL'):
            cur = 'pool'
        elif 'PLAYER ROSTER' in s:
            cur = 'roster'
        elif 'WEAPON USAGE' in s:
            cur = 'weapons'
        elif 'ATTACK TIMING' in s:
            cur = 'timing'
        elif 'SITE EXECUTION' in s:
            cur = 'sites'
        elif 'PLAYER OPENING ROUTES' in s:
            cur = 'routes'
        elif 'ROUND-BY-ROUND' in s:
            cur = 'rounds'
        elif 'COUNTER-STRAT BRIEF' in s:
            cur = 'brief'

        sections[cur].append(L)
        i += 1

    return sections


def build_docx(team: str, series_id=None, map_filter=None):
    # ── 1. Capture full text output from antistrat.py ─────────────────────────
    print("Running analysis (this takes 30-90s for the AI section)...")
    cmd = [sys.executable, str(Path(__file__).parent / 'antistrat.py'),
           '--team', team]
    if series_id:
        cmd += ['--series', str(series_id)]
    if map_filter:
        cmd += ['--map', map_filter]
    env = {'PYTHONIOENCODING': 'utf-8', 'PATH': __import__('os').environ['PATH']}

    result = subprocess.run(cmd, capture_output=True, text=True,
                            encoding='utf-8', env=env,
                            cwd=str(Path(__file__).parent))
    raw = result.stdout
    if not raw.strip():
        raw = result.stderr
        print("Warning: no stdout, using stderr")

    secs = _parse_sections(raw)

    # ── 2. Build document ─────────────────────────────────────────────────────
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── TITLE ─────────────────────────────────────────────────────────────────
    title_p = doc.add_heading(f'MATCH ANALYSIS — {team.upper()}', level=1)
    for r in title_p.runs:
        r.font.color.rgb = RED
        r.font.size      = Pt(20)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # subtitle from header lines
    for line in secs['header']:
        s = line.strip()
        if 'vs' in s and '|' in s:
            sp = _p(doc, s, color=GRAY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # ── MAP POOL ──────────────────────────────────────────────────────────────
    _h(doc, 'MAP POOL')
    _rule(doc)
    for line in secs['pool']:
        s = line.strip()
        if not s or 'MAP POOL' in s or s.startswith('=') or s.startswith('-'):
            continue
        _p(doc, s, size=10)
    doc.add_paragraph()

    # ── PLAYER ROSTER ─────────────────────────────────────────────────────────
    _h(doc, 'PLAYER ROSTER')
    _rule(doc)

    # Build table from roster lines
    data_rows = []
    for line in secs['roster']:
        s = line.strip()
        if not s or 'ROSTER' in s or s.startswith('=') or s.startswith('-') or s.startswith('PLAYER'):
            continue
        # line format: BABYBAY: 131 dmg/rd | FK 3% | FD 7% | 0 plants → SUPPORT / ANCHOR
        if ':' not in s:
            continue
        try:
            name, rest = s.split(':', 1)
            name = name.strip().lstrip('*').strip()
            parts = [x.strip() for x in rest.split('|')]
            dmg  = parts[0].split()[0] if parts else ''
            fk   = parts[1].replace('FK','').strip() if len(parts) > 1 else ''
            fd   = parts[2].replace('FD','').strip() if len(parts) > 2 else ''
            plt  = parts[3].split()[0] if len(parts) > 3 else ''
            role_raw = parts[4] if len(parts) > 4 else ''
            import re
            role = re.sub(r'(?:->|→).*', '', role_raw).strip()
            data_rows.append([name, dmg, fk, fd, plt, role])
        except (IndexError, ValueError):
            continue

    if data_rows:
        headers = ['PLAYER', 'DMG/RD', 'FK%', 'FD%', 'PLANTS', 'ROLE']
        tbl = doc.add_table(rows=1 + len(data_rows), cols=6)
        tbl.style = 'Table Grid'
        # header row
        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]
            c.text = h
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = Pt(9)
            c.paragraphs[0].runs[0].font.color.rgb = WHITE
            _shd(c, '1F497D')
        # data rows
        for ri, row_data in enumerate(data_rows):
            for ci, val in enumerate(row_data):
                c = tbl.rows[ri + 1].cells[ci]
                c.text = val
                c.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── WEAPONS ───────────────────────────────────────────────────────────────
    _h(doc, 'WEAPON USAGE')
    _rule(doc)
    for line in secs['weapons']:
        s = line.strip()
        if not s or 'WEAPON' in s or s.startswith('=') or s.startswith('-'):
            continue
        _p(doc, s, size=10)
    doc.add_paragraph()

    # ── TIMING ────────────────────────────────────────────────────────────────
    _h(doc, 'ATTACK TIMING')
    _rule(doc)
    for line in secs['timing']:
        s = line.strip()
        if not s or 'TIMING' in s or s.startswith('=') or s.startswith('-'):
            continue
        if s.startswith('->') or s.startswith('→'):
            _p(doc, s.lstrip('->→ '), bold=True, color=RED, size=10)
        elif s.startswith('IMPLICATION'):
            _p(doc, s, italic=True, color=GRAY, size=9)
        else:
            _p(doc, s, size=10)
    doc.add_paragraph()

    # ── SITES ─────────────────────────────────────────────────────────────────
    _h(doc, 'SITE EXECUTION TENDENCIES')
    _rule(doc)
    for line in secs['sites']:
        s = line.strip()
        if not s or 'SITE' in s or s.startswith('=') or s.startswith('-'):
            continue
        _p(doc, s, size=10)
    doc.add_paragraph()

    # ── ROUTES ────────────────────────────────────────────────────────────────
    _h(doc, 'PLAYER OPENING ROUTES')
    _p(doc, 'Position at ~10s into ATK rounds — use to predict setups',
       italic=True, color=GRAY, size=9)
    _rule(doc)
    for line in secs['routes']:
        s = line.strip()
        if not s or 'ROUTES' in s or s.startswith('=') or s.startswith('-'):
            continue
        if s.startswith('['):
            _p(doc, s, bold=True, color=BLUE, size=10, indent=0.1)
        else:
            _p(doc, s, size=10, indent=0.3)
    doc.add_paragraph()

    # ── ROUNDS ────────────────────────────────────────────────────────────────
    _h(doc, 'ROUND-BY-ROUND ATK SETUPS')
    _rule(doc)
    for line in secs['rounds']:
        s = line.strip()
        if not s or 'ROUND-BY-ROUND' in s or s.startswith('==='):
            continue
        if s.startswith('---') or s.startswith('─'):
            _rule(doc)
        elif s.startswith('[') and ']' in s and not s.startswith('[WIN]') and not s.startswith('[LOSS]'):
            _p(doc, s, bold=True, color=BLUE, size=11)
            _rule(doc)
        elif s.startswith('R') and '→' in s:
            won = '[WIN]' in s
            color = GREEN if won else RED
            _p(doc, s, bold=True, color=color, size=10)
        elif 'Opening setup' in s:
            _p(doc, s, bold=True, color=BLACK, size=9, indent=0.25)
        elif 'Late-round' in s:
            _p(doc, s, bold=True, color=GRAY, size=9, indent=0.25)
        elif '@' in s and 'OPEN' not in s:
            _p(doc, s, size=9, indent=0.5)
        else:
            _p(doc, s, size=9, indent=0.25)
    doc.add_paragraph()

    # ── COUNTER-STRAT BRIEF ───────────────────────────────────────────────────
    _h(doc, 'COUNTER-STRAT BRIEF', level=1, color=RED, size=16)
    doc.add_paragraph()

    in_section = False
    for line in secs['brief']:
        s = line.strip()
        if not s or 'COUNTER-STRAT BRIEF' in s:
            continue
        if s.startswith('===') or s.startswith('---') or s.startswith('─'):
            _rule(doc)
        elif ('PEARL CT' in s.upper() or 'PEARL ATK' in s.upper()
              or ' CT (' in s.upper() or ' ATK (' in s.upper()
              or s.upper().startswith('CT (') or s.upper().startswith('ATK (')):
            _h(doc, s, level=2, color=BLUE)
            _rule(doc)
            in_section = True
        elif s.startswith('[!]'):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(s[3:].strip())
            r.font.size = Pt(10)
            r.font.color.rgb = BLACK
        elif s.startswith('>>'):
            _p(doc, s, bold=True, color=RED, size=10, indent=0.1)
        else:
            _p(doc, s, size=10)

    # ── SAVE ──────────────────────────────────────────────────────────────────
    safe = team.replace(' ', '_').replace('/', '-')
    mtag = f'_{map_filter}' if map_filter else ''
    stag = f'_s{series_id}' if series_id else ''
    out  = Path(__file__).parent / f'antistrat_{safe}{mtag}{stag}.docx'
    doc.save(str(out))
    return out


# ── CLI ───────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    ap = argparse.ArgumentParser()
    ap.add_argument('--team',   required=True)
    ap.add_argument('--series', type=int, default=None)
    ap.add_argument('--map',    default=None)
    args = ap.parse_args()

    path = build_docx(args.team, args.series, args.map)
    print(f'\nSaved: {path}')
